import pandas as pd
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LinearRegression, Ridge
from sklearn.metrics import mean_squared_error, r2_score
import statsmodels.api as sm
from sklearn.preprocessing import PolynomialFeatures
from docx import Document




# 读取 Excel 文件的第二个表格
# df = pd.read_excel('指标化数据.xlsx', sheet_name=1)
df = pd.read_excel('模拟数据.xlsx')
# # 使用 map 方法将 1 替换为 0，2 替换为 1
# df["性别"] = df["性别"].map({1: 0, 2: 1})

# 选择需要标准化的列
cols_to_standardize = ["购买频率","性别", "年龄", "受教育水平", "职业", "月收入"]
# cols_to_standardize = ["性别", "职业", "月收入"]

# 初始化标准化处理器
scaler = StandardScaler()

# 对选择的列进行标准化
df[cols_to_standardize] = scaler.fit_transform(df[cols_to_standardize])


# print(df.head())

# 定义因变量和自变量
X = df[["品牌意识综合指标", "产品偏好综合指标", "消费心理综合指标", "性别",
	 "年龄", "受教育水平", "职业", "月收入"]]

y = df["购买频率"]

# 添加常量项以适配 statsmodels
X_sm = sm.add_constant(X)


#一般线性回归
# 线性回归模型
linear_model = sm.OLS(y, X_sm).fit()
print("线性回归模型报告：")
print(linear_model.summary())



# # 2. 拟合普通最小二乘（OLS）模型和稳健回归模型
# X_with_const = sm.add_constant(X)  # 添加常数项（截距）
# ols_model = sm.OLS(y, X_with_const).fit()
# rlm_model = sm.RLM(y, X_with_const, M=sm.robust.norms.HuberT()).fit()
# # rlm_model = sm.RLM(y, X_with_const, M=sm.robust.norms.AndrewWave()).fit()
# # 3. 生成模型报告
# print("普通最小二乘（OLS）模型报告")
# print(ols_model.summary())

# print("\n稳健回归模型报告")
# print(rlm_model.summary())






# # 创建 Word 文档
# doc = Document()

# # 添加标题
# doc.add_heading('线性回归分析报告', 0)

# # 添加回归模型的摘要
# doc.add_heading('模型摘要', level=1)
# doc.add_paragraph(str(linear_model.summary()))

# # 保存文档
# doc.save("regression_report.docx")





