import pandas as pd
from sklearn.tree import DecisionTreeRegressor, plot_tree
from sklearn.model_selection import train_test_split
from sklearn.metrics import mean_squared_error, r2_score
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

# 加载数据集
data = pd.read_excel('模拟数据.xlsx')
X = data[["品牌意识综合指标", "产品偏好综合指标", 
          "消费心理综合指标", "性别","年龄", "受教育水平", 
          "职业", "月收入"]]
y = data["购买频率"]

# 分割数据集
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.25, random_state=0)

# 构建决策树回归模型
regressor = DecisionTreeRegressor(
    max_depth=3, 
    min_samples_split=25, 
    min_samples_leaf=20, 
    random_state=42
)
regressor.fit(X_train, y_train)

# 预测和评估模型
y_pred = regressor.predict(X_test)
mse = mean_squared_error(y_test, y_pred)
rmse = np.sqrt(mse)
r2 = r2_score(y_test, y_pred)

# 输出评估报告
print("回归模型评估报告")
print("----------------------------")
print(f"均方误差 (MSE): {mse:.4f}")
print(f"均方根误差 (RMSE): {rmse:.4f}")
print(f"R² 分数: {r2:.4f}")

# 输出特征重要性
print("\n特征重要性:")
feature_importances = regressor.feature_importances_
for feature, importance in zip(X.columns, feature_importances):
    print(f"{feature}: {importance:.4f}")

# 可视化决策树
plt.figure(figsize=(16, 12))
plt.rcParams['font.family'] = ['Microsoft YaHei']  # 设置字体为微软雅黑
plot_tree(regressor, 
          feature_names=X.columns, 
          filled=True, 
          rounded=True, 
          fontsize=10, 
          precision=2, 
          proportion=True, 
          label='all', 
          impurity=False)
plt.title("决策树回归 - 饮食品牌IP联名产品的购买偏好及发展前景分析", fontsize=18)

# 保存图片
plt.savefig("decision_tree.png", dpi=300, bbox_inches='tight')
plt.show()

# 创建Word文档并插入图片
doc = Document()
doc.add_heading("决策树分析报告", level=1)
doc.add_paragraph("此决策树模型分析了饮食品牌IP联名产品的购买频率及其发展前景。以下是决策树图像及其模型的评估结果。")

doc.add_heading("评估报告", level=2)
doc.add_paragraph(f"均方误差 (MSE): {mse:.4f}")
doc.add_paragraph(f"均方根误差 (RMSE): {rmse:.4f}")
doc.add_paragraph(f"R² 分数: {r2:.4f}")

doc.add_heading("特征重要性", level=2)
for feature, importance in zip(X.columns, feature_importances):
    doc.add_paragraph(f"{feature}: {importance:.4f}")

# 插入图片
doc.add_heading("决策树图像", level=2)
doc.add_picture("decision_tree.png", width=Inches(6))

# 保存Word文档
doc.save("决策树分析报告.docx")
